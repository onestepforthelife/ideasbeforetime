"""
Simple LinkedIn posts extractor - just get all text
"""
from docx import Document
import os

doc_path = r"C:\Users\araag\Documents\Cloudfare\linkedin post back till 30 nov 2025.docx"
output_path = r"C:\Users\araag\Documents\Cloudfare\linkedin-posts-extracted.txt"

print("=" * 80)
print("📝 EXTRACTING LINKEDIN POSTS")
print("=" * 80)

try:
    doc = Document(doc_path)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("LINKEDIN POSTS - EXTRACTED FROM WORD FILE\n")
        f.write("=" * 80 + "\n\n")
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                f.write(text + "\n")
        
        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        f.write(text + "\n")
    
    print(f"✅ Extracted {len(doc.paragraphs)} paragraphs")
    print(f"✅ Saved to: {output_path}")
    print(f"\n🎯 Next: Open the file and I'll format it into blog posts!")
    
except Exception as e:
    print(f"❌ Error: {e}")
