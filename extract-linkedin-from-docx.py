"""
Extract LinkedIn posts from Word document and convert to blog format
"""
import os
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
from datetime import datetime

def extract_images_from_docx(doc_path):
    """Extract images from Word document"""
    import zipfile
    from PIL import Image
    import io
    
    images = []
    try:
        with zipfile.ZipFile(doc_path, 'r') as zip_ref:
            # Get all image files from word/media/
            image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
            
            for idx, img_file in enumerate(image_files):
                img_data = zip_ref.read(img_file)
                # Save image
                img_name = f"linkedin-post-image-{idx+1}.{img_file.split('.')[-1]}"
                img_path = os.path.join('Cloudfare', 'images', img_name)
                
                with open(img_path, 'wb') as f:
                    f.write(img_data)
                
                images.append(img_name)
                print(f"✅ Extracted image: {img_name}")
    except Exception as e:
        print(f"⚠️ Could not extract images: {e}")
    
    return images

def parse_linkedin_posts(doc_path):
    """Parse LinkedIn posts from Word document"""
    try:
        doc = Document(doc_path)
        posts = []
        current_post = None
        post_content = []
        image_counter = 0
        
        print(f"\n📄 Reading: {doc_path}")
        print(f"📊 Total paragraphs: {len(doc.paragraphs)}\n")
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Check if this is a date line (potential post separator)
            if re.match(r'\d{1,2}[/-]\d{1,2}[/-]\d{4}', text) or re.match(r'[A-Z][a-z]+ \d{1,2}, \d{4}', text):
                # Save previous post if exists
                if current_post and post_content:
                    current_post['content'] = '\n\n'.join(post_content)
                    posts.append(current_post)
                
                # Start new post
                current_post = {
                    'date': text,
                    'title': '',
                    'category': 'Industry Insights',
                    'content': ''
                }
                post_content = []
                print(f"📅 Found post date: {text}")
            
            # Check for images in paragraph
            elif para._element.xpath('.//pic:pic'):
                image_counter += 1
                post_content.append(f"[IMAGE: linkedin-post-image-{image_counter}]")
                print(f"🖼️ Found image #{image_counter}")
            
            # Regular content
            else:
                if current_post:
                    # First non-date line is likely the title/first line
                    if not current_post['title']:
                        current_post['title'] = text[:100] + ('...' if len(text) > 100 else '')
                    post_content.append(text)
        
        # Save last post
        if current_post and post_content:
            current_post['content'] = '\n\n'.join(post_content)
            posts.append(current_post)
        
        print(f"\n✅ Found {len(posts)} posts\n")
        return posts
    
    except Exception as e:
        print(f"❌ Error reading Word file: {e}")
        return []

def convert_date(date_str):
    """Convert various date formats to YYYY-MM-DD"""
    try:
        # Try different date formats
        for fmt in ['%m/%d/%Y', '%d/%m/%Y', '%B %d, %Y', '%b %d, %Y', '%Y-%m-%d']:
            try:
                dt = datetime.strptime(date_str, fmt)
                return dt.strftime('%Y-%m-%d')
            except:
                continue
        return '2025-11-30'  # Default
    except:
        return '2025-11-30'

def categorize_post(content):
    """Auto-categorize post based on content"""
    content_lower = content.lower()
    
    if any(word in content_lower for word in ['leader', 'team', 'manage', 'culture']):
        return 'Leadership'
    elif any(word in content_lower for word in ['chemical', 'industry', 'market', 'supply']):
        return 'Industry Insights'
    elif any(word in content_lower for word in ['ai', 'technology', 'digital', 'innovation']):
        return 'Technology'
    elif any(word in content_lower for word in ['career', 'job', 'interview', 'resume']):
        return 'Career Advice'
    elif any(word in content_lower for word in ['strategy', 'business', 'growth', 'revenue']):
        return 'Business Strategy'
    else:
        return 'Industry Insights'

def save_to_linkedin_posts_txt(posts):
    """Save posts to linkedin-posts.txt format"""
    output_path = 'Cloudfare/linkedin-posts.txt'
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("━" * 80 + "\n")
        f.write("📝 LINKEDIN POSTS - EXTRACTED FROM WORD FILE\n")
        f.write("━" * 80 + "\n\n")
        
        for idx, post in enumerate(posts, 1):
            date = convert_date(post['date'])
            category = categorize_post(post['content'])
            
            f.write("---POST---\n")
            f.write(f"TITLE: {post['title']}\n")
            f.write(f"DATE: {date}\n")
            f.write(f"CATEGORY: {category}\n\n")
            f.write(post['content'])
            f.write("\n\n")
            
            print(f"✅ Post {idx}: {post['title'][:50]}... ({date})")
    
    print(f"\n✅ Saved {len(posts)} posts to: {output_path}")
    return output_path

def main():
    doc_path = r"C:\Users\araag\Documents\Cloudfare\linkedin post back till 30 nov 2025.docx"
    
    print("=" * 80)
    print("📝 LINKEDIN POSTS EXTRACTOR")
    print("=" * 80)
    
    # Check if file exists
    if not os.path.exists(doc_path):
        print(f"❌ File not found: {doc_path}")
        return
    
    # Extract images
    print("\n🖼️ Extracting images...")
    images = extract_images_from_docx(doc_path)
    
    # Parse posts
    print("\n📝 Parsing posts...")
    posts = parse_linkedin_posts(doc_path)
    
    if not posts:
        print("❌ No posts found!")
        return
    
    # Save to linkedin-posts.txt
    print("\n💾 Saving posts...")
    output_file = save_to_linkedin_posts_txt(posts)
    
    print("\n" + "=" * 80)
    print("✅ EXTRACTION COMPLETE!")
    print("=" * 80)
    print(f"\n📊 Summary:")
    print(f"   • Posts extracted: {len(posts)}")
    print(f"   • Images extracted: {len(images)}")
    print(f"   • Output file: {output_file}")
    print(f"\n🎯 Next step: Tell Kiro 'add posts' to create blog pages!")

if __name__ == "__main__":
    main()
